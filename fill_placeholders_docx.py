# README:
#   python fill_placeholders_docx.py --docx "C:\\path\\to\\input.docx" --placeholders placeholders.json
#
# Notes:
#   - Requires: pip install python-docx
#   - Output: Capstone_Filled.docx (in current directory unless --out is set)

import argparse
import json
from pathlib import Path
from docx import Document


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_in_paragraph(paragraph, replacements: dict):
    if not paragraph.runs:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text
    for k, v in replacements.items():
        if k in new_text and v is not None:
            new_text = new_text.replace(k, v)
    if new_text != full_text:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", required=True, help="Path to input .docx")
    ap.add_argument("--placeholders", required=True, help="Path to placeholders.json")
    ap.add_argument("--out", default="Capstone_Filled.docx", help="Output .docx path")
    args = ap.parse_args()

    doc_path = Path(args.docx)
    placeholders_path = Path(args.placeholders)
    if not doc_path.exists():
        raise SystemExit(f"Docx not found: {doc_path}")
    if not placeholders_path.exists():
        raise SystemExit(f"placeholders.json not found: {placeholders_path}")

    replacements = json.loads(placeholders_path.read_text(encoding="utf-8"))

    doc = Document(str(doc_path))
    for p in iter_paragraphs(doc):
        replace_in_paragraph(p, replacements)

    out_path = Path(args.out)
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
