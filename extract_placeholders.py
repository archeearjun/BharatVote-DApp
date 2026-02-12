# README:
#   python extract_placeholders.py --docx "C:\\path\\to\\BharatVote_Capstone_Submission.docx"
#
# Notes:
#   - Requires: pip install python-docx
#   - Output: placeholders.json (in current directory)

import argparse
import json
import re
from pathlib import Path
from docx import Document

TOKEN_RE = re.compile(r"<<MANUAL_[^>]+>>")


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def extract_tokens(doc_path: Path):
    doc = Document(str(doc_path))
    tokens = set()
    for p in iter_paragraphs(doc):
        text = p.text
        for m in TOKEN_RE.findall(text):
            tokens.add(m)
    return sorted(tokens)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", required=True, help="Path to the .docx file")
    ap.add_argument("--out", default="placeholders.json", help="Output JSON path")
    args = ap.parse_args()

    doc_path = Path(args.docx)
    if not doc_path.exists():
        raise SystemExit(f"Docx not found: {doc_path}")

    tokens = extract_tokens(doc_path)
    payload = {t: "" for t in tokens}

    out_path = Path(args.out)
    out_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    print(f"Wrote {len(tokens)} tokens to {out_path}")


if __name__ == "__main__":
    main()
